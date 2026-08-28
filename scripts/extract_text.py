"""
Extract raw text from a deal's source documents (PDF, DOCX) in data/raw/<deal>/
into plain-text files in data/working/<deal>/, one .txt per source file.

Usage:
    python scripts/extract_text.py <deal-slug>
"""

import sys
from pathlib import Path

import pymupdf
import docx
import pytesseract

ROOT = Path(__file__).resolve().parent.parent

TESSERACT_EXE = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
if TESSERACT_EXE.exists():
    pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_EXE)


def ocr_page(page: "pymupdf.Page") -> str:
    pix = page.get_pixmap(dpi=300)
    img = pymupdf.Pixmap(pix, 0) if pix.alpha else pix
    from PIL import Image
    import io

    image = Image.open(io.BytesIO(img.tobytes("png")))
    return pytesseract.image_to_string(image)


def extract_pdf(path: Path) -> str:
    parts = []
    with pymupdf.open(path) as pdf:
        for i, page in enumerate(pdf, start=1):
            parts.append(f"\n--- page {i} ---\n")
            text = page.get_text()
            if not text.strip():
                text = ocr_page(page)
                if text.strip():
                    parts.append("[OCR]\n")
            parts.append(text)
    return "".join(parts)


def extract_docx(path: Path) -> str:
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def main():
    if len(sys.argv) != 2:
        print("Usage: python scripts/extract_text.py <deal-slug>")
        sys.exit(1)

    deal = sys.argv[1]
    raw_dir = ROOT / "data" / "raw" / deal
    out_dir = ROOT / "data" / "working" / deal
    out_dir.mkdir(parents=True, exist_ok=True)

    if not raw_dir.exists():
        print(f"No such folder: {raw_dir}")
        sys.exit(1)

    for src in sorted(raw_dir.iterdir()):
        suffix = src.suffix.lower()
        if suffix == ".pdf":
            text = extract_pdf(src)
        elif suffix == ".docx":
            text = extract_docx(src)
        elif suffix in (".md", ".txt"):
            text = src.read_text(encoding="utf-8")
        else:
            print(f"skip (not text-extractable): {src.name}")
            continue

        out_path = out_dir / (src.stem + ".txt")
        out_path.write_text(text, encoding="utf-8")
        print(f"extracted: {src.name} -> {out_path.relative_to(ROOT)} ({len(text):,} chars)")


if __name__ == "__main__":
    main()
